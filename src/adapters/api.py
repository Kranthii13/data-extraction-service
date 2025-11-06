# src/infrastructure/api.py
from fastapi import FastAPI, UploadFile, File, Depends, HTTPException, Query, BackgroundTasks
from sqlalchemy.orm import Session
from typing import List, Optional, Dict, Any
import os
import uuid
import json
import logging
import time
from datetime import datetime
from enum import Enum

from src.core.models import Document, ExtractedData
from src.services.ports import IExtractionService
from src.adapters.dependencies import get_extraction_service, get_db
from src.config.app_config import config

def _limit_table_rows(table_rows: List[List[str]], max_rows: int = None) -> tuple[List[List[str]], dict]:
    """
    Limit table rows to prevent browser crashes and return metadata about truncation.
    
    Returns:
        tuple: (limited_rows, metadata)
    """
    if max_rows is None:
        max_rows = config.large_file.max_storage_rows
    
    if not table_rows or len(table_rows) <= max_rows:
        return table_rows, {
            'is_truncated': False,
            'original_row_count': len(table_rows) if table_rows else 0,
            'stored_row_count': len(table_rows) if table_rows else 0
        }
    
    # Truncate rows
    limited_rows = table_rows[:max_rows]
    metadata = {
        'is_truncated': True,
        'original_row_count': len(table_rows),
        'stored_row_count': len(limited_rows),
        'truncation_reason': 'Large table truncated to prevent browser crashes'
    }
    
    logger.warning(f"Large table detected ({len(table_rows)} rows). Truncated to {max_rows} rows to prevent memory issues.")
    
    return limited_rows, metadata

logger = logging.getLogger(__name__)

# Check if production dependencies are available
USE_CELERY = False
try:
    import celery
    import redis
    
    # Test Redis connection
    redis_client = redis.Redis(
        host=os.getenv('REDIS_HOST', 'localhost'),
        port=int(os.getenv('REDIS_PORT', 6379)),
        db=int(os.getenv('REDIS_DB', 0)),
        decode_responses=True
    )
    redis_client.ping()
    
    # Setup Celery
    celery_app = celery.Celery(
        'document_processor',
        broker=os.getenv('CELERY_BROKER_URL', 'redis://localhost:6379/0'),
        backend=os.getenv('CELERY_RESULT_BACKEND', 'redis://localhost:6379/0')
    )
    
    @celery_app.task
    def process_document_task(document_data: dict) -> dict:
        """Celery task for document processing."""
        import base64
        import time
        import hashlib
        from src.core.models import Document
        from src.adapters.dependencies import get_extraction_service, SessionLocal
        from src.services.tabular_processor import TabularProcessor
        from src.adapters.database.models import DocumentRecord
        
        content = base64.b64decode(document_data['content'])
        document = Document(content=content, filename=document_data['filename'])
        
        db = SessionLocal()
        try:
            # Check if this is a tabular file and handle as table
            if _is_tabular_file(document.filename, content):
                # Process as tabular data
                start_time = time.time()
                
                # Detect file type
                file_type = TabularProcessor.detect_file_type(document.filename, content)
                if not file_type:
                    raise ValueError("Unable to detect tabular file type")
                
                # Load as DataFrame
                df = TabularProcessor.load_dataframe(content, file_type, document.filename)
                
                # Create table data structure
                table_data = TabularProcessor.create_table_data(df, file_type, document.filename)
                
                # Store in database
                file_hash = hashlib.sha256(content).hexdigest()
                
                # Check for existing document
                existing_doc = db.query(DocumentRecord).filter(DocumentRecord.file_hash == file_hash).first()
                if existing_doc:
                    document_id = existing_doc.id
                    action = "duplicate"
                else:
                    # Create new document record
                    db_document = DocumentRecord(
                        filename=document.filename,
                        file_extension=f".{file_type}",
                        file_size=len(content),
                        file_hash=file_hash,
                        full_text="",  # Tabular files don't have full text
                        page_count=1,
                        word_count=len(df) * len(df.columns),
                        processing_method=f"tabular_{file_type}",
                        has_ocr_content=0,  # Convert boolean False to integer 0
                        tables_data=table_data,
                        table_count=1
                    )
                    
                    db.add(db_document)
                    db.commit()
                    db.refresh(db_document)
                    document_id = db_document.id
                    action = "created"
                
                processing_time = int((time.time() - start_time) * 1000)
                
                # Create result with size limits for Celery task
                celery_result = {
                    "id": document_id,
                    "filename": document.filename,
                    "action": action,
                    "processing_time_ms": processing_time,
                    "file_size_bytes": len(content),
                    "data_format": "table",
                    "table_preview": TabularProcessor.get_preview_data(df),  # Already limited
                    "table_info": {
                        "shape": f"{len(df)} rows × {len(df.columns)} columns",
                        "columns": list(df.columns),
                        "data_types": {col: str(df[col].dtype) for col in df.columns}
                    },
                    "data_quality": TabularProcessor.analyze_data_quality(df)
                }
                
                # Apply size limits to prevent browser crashes in async responses
                return _apply_size_limits_to_task_result(celery_result)
            else:
                # Regular document processing
                service = get_extraction_service(db)
                result = service.extract_from_document(document)
                
                # Create result for Celery task with size limits
                celery_result = {
                    'id': getattr(result, 'id', None),
                    'filename': document_data['filename'],
                    'page_count': result.page_count,
                    'processing_method': result.processing_method,
                    'has_ocr_content': result.has_ocr_content,
                    'text_preview': result.full_text[:200] + "..." if len(result.full_text) > 200 else result.full_text,
                    'table_count': result.table_count
                }
                
                # Include limited table data if present
                if hasattr(result, 'tables') and result.tables:
                    limited_tables = []
                    for table in result.tables:
                        table_dict = {
                            "table_index": table.table_index,
                            "page_number": table.page_number,
                            "title": table.title,
                            "row_count": table.row_count,
                            "column_count": table.column_count,
                            "table_type": table.table_type,
                            "confidence_score": table.confidence_score,
                            "extraction_method": table.extraction_method,
                            # Include truncation metadata
                            "is_truncated": getattr(table, 'is_truncated', False),
                            "original_row_count": getattr(table, 'original_row_count', table.row_count),
                            "stored_row_count": getattr(table, 'stored_row_count', table.row_count)
                        }
                        
                        # Add limited row data (preview only for Celery results)
                        if table.rows:
                            from src.config.app_config import config
                            preview_size = min(config.large_file.max_response_rows, len(table.rows))
                            table_dict['rows_preview'] = table.rows[:preview_size]
                            table_dict['preview_truncated'] = len(table.rows) > preview_size
                            table_dict['total_rows_available'] = len(table.rows)
                        
                        limited_tables.append(table_dict)
                    
                    celery_result['tables'] = limited_tables
                
                # Apply size limits to prevent browser crashes in async responses
                return _apply_size_limits_to_task_result(celery_result)
        finally:
            db.close()
    
    USE_CELERY = True
    logger.info("Using Celery + Redis for task processing")
    
except Exception as e:
    logger.info(f"Using BackgroundTasks for task processing: {e}")
    # Simple task store for development
    task_store: Dict[str, Dict[str, Any]] = {}

class TaskStatus(str, Enum):
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"

app = FastAPI(
    title="Data Extraction Service",
    description="Document processing with async support",
    version="1.0.0"
)

# Task management functions
def store_task_status(task_id: str, status: dict):
    """Store task status."""
    if USE_CELERY:
        redis_client.setex(f"task:{task_id}", 3600, json.dumps(status))
    else:
        task_store[task_id] = status

def get_task_status(task_id: str) -> Optional[dict]:
    """Get task status."""
    if USE_CELERY:
        data = redis_client.get(f"task:{task_id}")
        return json.loads(data) if data else None
    else:
        return task_store.get(task_id)

def update_task_status(task_id: str, updates: dict):
    """Update task status."""
    if USE_CELERY:
        current = get_task_status(task_id) or {}
        current.update(updates)
        redis_client.setex(f"task:{task_id}", 3600, json.dumps(current))
    else:
        if task_id in task_store:
            task_store[task_id].update(updates)

# Background processing for development
async def process_document_background(task_id: str, document: Document, db: Session):
    """Process document in background."""
    try:
        update_task_status(task_id, {"status": TaskStatus.PROCESSING})
        
        # Check if this is a tabular file and handle as table
        if _is_tabular_file(document.filename, document.content):
            # Process as tabular data
            import time
            from src.services.tabular_processor import TabularProcessor
            import hashlib
            from src.adapters.database.models import DocumentRecord
            
            start_time = time.time()
            
            # Detect file type
            file_type = TabularProcessor.detect_file_type(document.filename, document.content)
            if not file_type:
                raise ValueError("Unable to detect tabular file type")
            
            # Load as DataFrame
            df = TabularProcessor.load_dataframe(document.content, file_type, document.filename)
            
            # Create table data structure
            table_data = TabularProcessor.create_table_data(df, file_type, document.filename)
            
            # Store in database
            file_hash = hashlib.sha256(document.content).hexdigest()
            
            # Check for existing document
            existing_doc = db.query(DocumentRecord).filter(DocumentRecord.file_hash == file_hash).first()
            if existing_doc:
                document_id = existing_doc.id
                action = "duplicate"
            else:
                # Create new document record
                db_document = DocumentRecord(
                    filename=document.filename,
                    file_extension=f".{file_type}",
                    file_size=len(document.content),
                    file_hash=file_hash,
                    full_text="",  # Tabular files don't have full text
                    page_count=1,
                    word_count=len(df) * len(df.columns),
                    processing_method=f"tabular_{file_type}",
                    has_ocr_content=0,  # Convert boolean False to integer 0
                    tables_data=table_data,
                    table_count=1
                )
                
                db.add(db_document)
                db.commit()
                db.refresh(db_document)
                document_id = db_document.id
                action = "created"
            
            processing_time = int((time.time() - start_time) * 1000)
            
            # Create result with size limits applied
            task_result = {
                "id": document_id,
                "filename": document.filename,
                "action": action,
                "processing_time_ms": processing_time,
                "file_size_bytes": len(document.content),
                "data_format": "table",
                "table_preview": TabularProcessor.get_preview_data(df),  # Already limited by get_preview_data
                "table_info": {
                    "shape": f"{len(df)} rows × {len(df.columns)} columns",
                    "columns": list(df.columns),
                    "data_types": {col: str(df[col].dtype) for col in df.columns}
                },
                "data_quality": TabularProcessor.analyze_data_quality(df)
            }
            
            # Apply additional size limits to prevent browser crashes
            limited_result = _apply_size_limits_to_task_result(task_result)
            
            update_task_status(task_id, {
                "status": TaskStatus.COMPLETED,
                "result": limited_result
            })
        else:
            # Regular document processing
            service = get_extraction_service(db)
            result = service.extract_from_document(document)
            
            # Create result for regular documents
            task_result = {
                "id": getattr(result, 'id', None),
                "filename": document.filename,
                "page_count": result.page_count,
                "processing_method": result.processing_method,
                "has_ocr_content": result.has_ocr_content,
                "text_preview": result.full_text[:200] + "..." if len(result.full_text) > 200 else result.full_text,
                "table_count": result.table_count
            }
            
            # Include limited table data if present
            if hasattr(result, 'tables') and result.tables:
                # Convert tables to serializable format with size limits
                limited_tables = []
                for table in result.tables:
                    table_dict = {
                        "table_index": table.table_index,
                        "page_number": table.page_number,
                        "title": table.title,
                        "row_count": table.row_count,
                        "column_count": table.column_count,
                        "table_type": table.table_type,
                        "confidence_score": table.confidence_score,
                        "extraction_method": table.extraction_method,
                        # Include truncation metadata
                        "is_truncated": getattr(table, 'is_truncated', False),
                        "original_row_count": getattr(table, 'original_row_count', table.row_count),
                        "stored_row_count": getattr(table, 'stored_row_count', table.row_count)
                    }
                    
                    # Add limited row data (preview only for async results)
                    if table.rows:
                        preview_size = min(config.large_file.max_response_rows, len(table.rows))
                        table_dict['rows_preview'] = table.rows[:preview_size]
                        table_dict['preview_truncated'] = len(table.rows) > preview_size
                        table_dict['total_rows_available'] = len(table.rows)
                    
                    limited_tables.append(table_dict)
                
                task_result['tables'] = limited_tables
            
            # Apply additional size limits
            limited_result = _apply_size_limits_to_task_result(task_result)
            
            update_task_status(task_id, {
                "status": TaskStatus.COMPLETED,
                "result": limited_result
            })
        
    except Exception as e:
        logger.error(f"Background processing failed: {e}")
        update_task_status(task_id, {
            "status": TaskStatus.FAILED,
            "error": str(e)
        })

# API Endpoints
@app.post("/extract/")
async def extract_sync(file: UploadFile = File(...), db: Session = Depends(get_db)):
    """Process document synchronously with performance tracking and CSV table handling."""
    import time
    start_time = time.time()
    
    content = await file.read()
    
    # Check if this is a tabular file and handle as table
    if _is_tabular_file(file.filename, content):
        return await _process_tabular_as_table(file, content, start_time, db)
    
    # Regular document processing
    document = Document(content=content, filename=file.filename)
    
    service = get_extraction_service(db)
    
    # Get the action info by checking if document exists first
    import hashlib
    from src.adapters.database.models import DocumentRecord
    
    file_hash = hashlib.sha256(content).hexdigest()
    existing = db.query(DocumentRecord).filter(DocumentRecord.file_hash == file_hash).first()
    action = "updated" if existing else "created"
    
    # Process the document
    result = service.extract_from_document(document)
    
    # Calculate processing time
    processing_time = round((time.time() - start_time) * 1000)  # milliseconds
    
    # Convert tables to serializable format with size limits for ALL document types
    tables_data = []
    if result.tables:
        for table in result.tables:
            # Apply response size limits to prevent browser crashes
            limited_rows = table.rows
            response_truncated = False
            
            if table.rows and len(table.rows) > config.large_file.max_response_rows:
                limited_rows = table.rows[:config.large_file.max_response_rows]
                response_truncated = True
                logger.info(f"Sync response truncated: showing {config.large_file.max_response_rows} of {len(table.rows)} rows for table {table.table_index}")
            
            # Create semantic data format from headers and limited rows
            data_records = []
            if table.headers and limited_rows:
                for row in limited_rows:
                    if len(row) == len(table.headers):
                        record = {header: (value if value is not None else None) for header, value in zip(table.headers, row)}
                        data_records.append(record)
            
            table_dict = {
                "table_index": table.table_index,
                "page_number": table.page_number,
                "title": table.title,
                "context_before": table.context_before,
                "context_after": table.context_after,
                "section_heading": table.section_heading,
                "headers": table.headers,  # Include headers for compatibility
                "data": data_records,  # Key-value format with size limits
                "row_count": table.row_count,
                "column_count": table.column_count,
                "table_type": table.table_type,
                "confidence_score": table.confidence_score,
                "extraction_method": table.extraction_method,
                # Add response truncation metadata
                "response_truncated": response_truncated,
                "response_sample_size": len(data_records),
                "total_rows_available": table.row_count,
                # Include storage truncation info if available
                "storage_truncated": getattr(table, 'is_truncated', False),
                "storage_truncation_reason": getattr(table, 'truncation_reason', None),
                "original_row_count": getattr(table, 'original_row_count', table.row_count),
                "stored_row_count": getattr(table, 'stored_row_count', table.row_count)
            }
            tables_data.append(table_dict)
    
    # Add action information to response
    response_data = {
        "id": result.id,
        "filename": result.filename,
        "full_text": result.full_text,
        "page_count": result.page_count,
        "has_ocr_content": result.has_ocr_content,
        "processing_method": result.processing_method,
        "table_count": result.table_count,
        "tables": tables_data,
        "action": action,
        "processing_time_ms": processing_time,
        "file_size_bytes": len(content)
    }
    
    return response_data

@app.post("/extract/async/")
async def extract_async(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """Process document asynchronously."""
    content = await file.read()
    document = Document(content=content, filename=file.filename)
    task_id = str(uuid.uuid4())
    
    # Store initial task status
    store_task_status(task_id, {
        "task_id": task_id,
        "status": TaskStatus.PENDING,
        "filename": file.filename,
        "created_at": datetime.now().isoformat()
    })
    
    if USE_CELERY:
        # Use Celery for production
        import base64
        task_data = {
            'content': base64.b64encode(content).decode('utf-8'),
            'filename': file.filename
        }
        celery_task = process_document_task.delay(task_data)
        update_task_status(task_id, {"celery_task_id": celery_task.id})
    else:
        # Use BackgroundTasks for development
        background_tasks.add_task(process_document_background, task_id, document, db)
    
    return {
        "task_id": task_id,
        "status": "pending"
    }

@app.get("/extract/status/{task_id}")
async def get_status(task_id: str):
    """Get task status with size-limited results to prevent browser crashes."""
    status = get_task_status(task_id)
    if not status:
        raise HTTPException(status_code=404, detail="Task not found")
    
    # Check Celery status if using Celery
    if USE_CELERY and 'celery_task_id' in status:
        celery_task = celery_app.AsyncResult(status['celery_task_id'])
        if celery_task.state == 'SUCCESS':
            # Apply size limits to Celery result before storing
            celery_result = celery_task.result
            if isinstance(celery_result, dict):
                celery_result = _apply_size_limits_to_task_result(celery_result)
            
            update_task_status(task_id, {
                "status": TaskStatus.COMPLETED,
                "result": celery_result
            })
            status = get_task_status(task_id)
        elif celery_task.state == 'FAILURE':
            update_task_status(task_id, {
                "status": TaskStatus.FAILED,
                "error": str(celery_task.info)
            })
            status = get_task_status(task_id)
    
    # Apply size limits to the result before returning (for both Celery and BackgroundTasks)
    if status and status.get('result'):
        status['result'] = _apply_size_limits_to_task_result(status['result'])
    
    return status

@app.get("/documents/{document_id}")
async def get_document(document_id: int, db: Session = Depends(get_db)):
    """Get document by ID with limited table data to prevent browser crashes."""
    service = get_extraction_service(db)
    document = service.get_document_by_id(document_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Convert to dict and limit table data for response
    doc_dict = document.dict()
    
    # Limit table data to prevent browser crashes (applies to all document types)
    if doc_dict.get('tables'):
        limited_tables = []
        for table in doc_dict['tables']:
            # Limit rows to prevent browser crashes for ALL document types
            if table.get('rows') and len(table['rows']) > config.large_file.max_response_rows:
                original_rows = len(table['rows'])
                table['rows'] = table['rows'][:config.large_file.max_response_rows]
                table['response_truncated'] = True
                table['response_sample_size'] = config.large_file.max_response_rows
                table['total_rows_available'] = table.get('original_row_count', original_rows)
                logger.info(f"API response truncated: showing {config.large_file.max_response_rows} of {original_rows} rows for table {table.get('table_index', 'unknown')}")
            else:
                table['response_truncated'] = False
                table['response_sample_size'] = len(table.get('rows', []))
            
            # Include storage truncation info if available
            if table.get('is_truncated'):
                table['storage_truncated'] = True
                table['storage_truncation_reason'] = table.get('truncation_reason')
            else:
                table['storage_truncated'] = False
                
            limited_tables.append(table)
        doc_dict['tables'] = limited_tables
    
    return doc_dict

@app.get("/documents/", response_model=List[ExtractedData])
async def get_documents(
    limit: int = Query(100, ge=1, le=1000),
    offset: int = Query(0, ge=0),
    db: Session = Depends(get_db)
):
    """Get all documents."""
    service = get_extraction_service(db)
    return service.get_all_documents(limit=limit, offset=offset)

@app.get("/search/", response_model=List[ExtractedData])
async def search_documents(
    q: str = Query(..., min_length=1),
    limit: int = Query(100, ge=1, le=1000),
    db: Session = Depends(get_db)
):
    """Search documents."""
    service = get_extraction_service(db)
    return service.search_documents(search_term=q, limit=limit)

@app.get("/documents/{document_id}/tables")
async def get_document_tables(document_id: int, db: Session = Depends(get_db)):
    """Get all tables from a specific document."""
    service = get_extraction_service(db)
    document = service.get_document_by_id(document_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Enhance response with contextual information
    tables_with_context = []
    for table in document.tables:
        table_info = {
            "table_index": table.table_index,
            "page_number": table.page_number,
            "title": table.title,
            "context_before": table.context_before,
            "context_after": table.context_after,
            "section_heading": table.section_heading,
            "table_type": table.table_type,
            "confidence_score": table.confidence_score,
            "headers": table.headers,
            "row_count": table.row_count,
            "column_count": table.column_count,
            "rows": table.rows[:5] if table.rows else [],  # Show first 5 rows as preview
            "total_rows": len(table.rows) if table.rows else 0
        }
        tables_with_context.append(table_info)
    
    return {
        "document_id": document_id,
        "filename": document.filename,
        "table_count": document.table_count,
        "tables": tables_with_context
    }

@app.get("/documents/{document_id}/tables/{table_index}")
async def get_document_table(
    document_id: int, 
    table_index: int, 
    format: str = Query("json", regex="^(json|html|markdown|context)$"),
    page: int = Query(1, ge=1),
    page_size: int = Query(100, ge=1, le=1000),
    db: Session = Depends(get_db)
):
    """Get a specific table from a document with pagination to prevent browser crashes."""
    from src.adapters.database.models import DocumentRecord
    
    # Get document with tables
    document = db.query(DocumentRecord).filter(DocumentRecord.id == document_id).first()
    if not document or not document.tables_data:
        raise HTTPException(status_code=404, detail="Document or table not found")
    
    # Find the specific table
    table_data = None
    for table in document.tables_data:
        if table.get("table_index") == table_index:
            table_data = table
            break
    
    if not table_data:
        raise HTTPException(status_code=404, detail="Table not found")
    
    # Apply pagination to prevent browser crashes
    def paginate_data(data_list, page_num, size):
        if not data_list:
            return []
        start_idx = (page_num - 1) * size
        end_idx = start_idx + size
        return data_list[start_idx:end_idx]
    
    if format == "html":
        return {"table_html": table_data.get("table_html")}
    elif format == "markdown":
        return {"table_markdown": table_data.get("table_markdown")}
    elif format == "context":
        # For context, limit the data to prevent crashes (works for all document types)
        rows = table_data.get("rows", [])
        data = table_data.get("data", [])
        
        # Use data field if available (for CSV files and new format), otherwise use rows (for PDF/DOCX/HTML)
        if data and len(data) > 0:
            paginated_data = paginate_data(data, page, page_size)
            total_rows = len(data)
            data_format = "key_value"  # Modern format
        elif rows and len(rows) > 0:
            # Convert rows to key-value format for consistency across all document types
            headers = table_data.get("headers", [])
            if headers:
                rows_as_records = []
                for row in rows:
                    if len(row) == len(headers):
                        record = {header: (value if value is not None else None) for header, value in zip(headers, row)}
                        rows_as_records.append(record)
                paginated_data = paginate_data(rows_as_records, page, page_size)
                total_rows = len(rows_as_records)
                data_format = "converted_rows"
            else:
                # Fallback to raw rows if no headers
                paginated_data = paginate_data(rows, page, page_size)
                total_rows = len(rows)
                data_format = "raw_rows"
        else:
            paginated_data = []
            total_rows = 0
            data_format = "empty"
        
        return {
            "table_index": table_data.get("table_index"),
            "page_number": table_data.get("page_number"),
            "title": table_data.get("title"),
            "context_before": table_data.get("context_before"),
            "context_after": table_data.get("context_after"),
            "section_heading": table_data.get("section_heading"),
            "table_type": table_data.get("table_type"),
            "confidence_score": table_data.get("confidence_score"),
            "headers": table_data.get("headers"),
            "data": paginated_data,  # Use paginated data
            "row_count": table_data.get("row_count"),
            "column_count": table_data.get("column_count"),
            "pagination": {
                "page": page,
                "page_size": page_size,
                "total_rows": total_rows,
                "total_pages": (total_rows + page_size - 1) // page_size if total_rows > 0 else 0,
                "has_next": page * page_size < total_rows,
                "has_prev": page > 1
            },
            "data_format": data_format,
            "supports_all_document_types": True,  # Works for PDF, DOCX, HTML, CSV, etc.
            "truncation_info": {
                "storage_truncated": table_data.get("is_truncated", False),
                "storage_reason": table_data.get("truncation_reason"),
                "original_row_count": table_data.get("original_row_count"),
                "stored_row_count": table_data.get("stored_row_count")
            }
        }
    else:  # json (basic)
        # For basic JSON, also apply pagination (works for all document types)
        rows = table_data.get("rows", [])
        data = table_data.get("data", [])
        
        # Use data field if available (for CSV files and new format), otherwise use rows (for PDF/DOCX/HTML)
        if data and len(data) > 0:
            paginated_data = paginate_data(data, page, page_size)
            total_rows = len(data)
            data_format = "key_value"  # Modern format
        elif rows and len(rows) > 0:
            # Convert rows to key-value format for consistency across all document types
            headers = table_data.get("headers", [])
            if headers:
                rows_as_records = []
                for row in rows:
                    if len(row) == len(headers):
                        record = {header: (value if value is not None else None) for header, value in zip(headers, row)}
                        rows_as_records.append(record)
                paginated_data = paginate_data(rows_as_records, page, page_size)
                total_rows = len(rows_as_records)
                data_format = "converted_rows"
            else:
                # Fallback to raw rows if no headers
                paginated_data = paginate_data(rows, page, page_size)
                total_rows = len(rows)
                data_format = "raw_rows"
        else:
            paginated_data = []
            total_rows = 0
            data_format = "empty"
        
        return {
            "table_index": table_data.get("table_index"),
            "page_number": table_data.get("page_number"),
            "headers": table_data.get("headers"),
            "data": paginated_data,  # Use paginated data instead of full rows
            "row_count": table_data.get("row_count"),
            "column_count": table_data.get("column_count"),
            "pagination": {
                "page": page,
                "page_size": page_size,
                "total_rows": total_rows,
                "total_pages": (total_rows + page_size - 1) // page_size if total_rows > 0 else 0,
                "has_next": page * page_size < total_rows,
                "has_prev": page > 1
            },
            "data_format": data_format,
            "supports_all_document_types": True,  # Works for PDF, DOCX, HTML, CSV, etc.
            "truncation_info": {
                "storage_truncated": table_data.get("is_truncated", False),
                "storage_reason": table_data.get("truncation_reason"),
                "original_row_count": table_data.get("original_row_count"),
                "stored_row_count": table_data.get("stored_row_count")
            }
        }

@app.get("/tables/search")
async def search_tables(
    q: str = Query(..., min_length=1),
    limit: int = Query(100, ge=1, le=1000),
    db: Session = Depends(get_db)
):
    """Search within table content using PostgreSQL JSON queries."""
    from src.adapters.database.models import DocumentRecord
    from sqlalchemy import text
    
    # Search in tables_data JSON using PostgreSQL JSON operators
    query = text("""
        SELECT d.id, d.filename, d.tables_data, d.table_count
        FROM documents d
        WHERE d.tables_data IS NOT NULL 
        AND d.tables_data::text ILIKE :search_term
        LIMIT :limit
    """)
    
    results = db.execute(query, {
        "search_term": f"%{q}%",
        "limit": limit
    }).fetchall()
    
    # Process results to extract matching tables
    table_results = []
    for row in results:
        if row.tables_data:
            for table in row.tables_data:
                # Check if this table contains the search term
                table_text = table.get("table_text", "")
                if q.lower() in table_text.lower():
                    table_results.append({
                        "document_id": row.id,
                        "filename": row.filename,
                        "table_index": table.get("table_index"),
                        "page_number": table.get("page_number"),
                        "headers": table.get("headers"),
                        "row_count": table.get("row_count"),
                        "column_count": table.get("column_count"),
                        "table_text": table_text[:200] + "..." if len(table_text) > 200 else table_text
                    })
    
    return {
        "query": q,
        "total_results": len(table_results),
        "tables": table_results[:limit]
    }







@app.get("/tables/export/{document_id}/{table_index}")
async def export_table(
    document_id: int,
    table_index: int,
    format: str = Query("csv", regex="^(csv|excel|json)$"),
    db: Session = Depends(get_db)
):
    """Export a specific table in various formats."""
    from src.adapters.database.models import DocumentRecord
    from fastapi.responses import Response
    
    # Get document with tables
    document = db.query(DocumentRecord).filter(DocumentRecord.id == document_id).first()
    if not document or not document.tables_data:
        raise HTTPException(status_code=404, detail="Document or table not found")
    
    # Find the specific table
    table_data = None
    for table in document.tables_data:
        if table.get("table_index") == table_index:
            table_data = table
            break
    
    if not table_data:
        raise HTTPException(status_code=404, detail="Table not found")
    
    filename = f"table_{document_id}_{table_index}"
    
    if format == "csv":
        content = table_data.get("table_csv", "")
        media_type = "text/csv"
        filename += ".csv"
    elif format == "excel":
        # Create Excel file (would need openpyxl library)
        content = table_data.get("table_csv", "")  # Fallback to CSV for now
        media_type = "application/vnd.ms-excel"
        filename += ".xlsx"
    else:  # json
        import json
        table_json = {
            "headers": table_data.get("headers"),
            "rows": table_data.get("rows"),
            "metadata": {
                "table_type": table_data.get("table_type"),
                "column_types": table_data.get("column_types"),
                "confidence_score": table_data.get("confidence_score")
            }
        }
        content = json.dumps(table_json, indent=2)
        media_type = "application/json"
        filename += ".json"
    
    return Response(
        content=content,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@app.get("/tables/stats")
async def get_table_statistics(db: Session = Depends(get_db)):
    """Get comprehensive table extraction statistics."""
    from src.adapters.database.models import DocumentRecord
    from sqlalchemy import text, func
    
    # Basic document and table statistics
    doc_stats = db.query(
        func.count(DocumentRecord.id).label('total_documents'),
        func.sum(DocumentRecord.table_count).label('total_tables'),
        func.avg(DocumentRecord.table_count).label('avg_tables_per_doc'),
        func.count().filter(DocumentRecord.table_count > 0).label('docs_with_tables')
    ).first()
    
    # Table type distribution
    table_type_stats = db.execute(text("""
        SELECT 
            json_array_elements(tables_data)->>'table_type' as table_type,
            COUNT(*) as count
        FROM documents 
        WHERE tables_data IS NOT NULL
        GROUP BY json_array_elements(tables_data)->>'table_type'
        ORDER BY count DESC
    """)).fetchall()
    
    # Extraction method statistics
    extraction_method_stats = db.execute(text("""
        SELECT 
            json_array_elements(tables_data)->>'extraction_method' as method,
            COUNT(*) as count,
            AVG((json_array_elements(tables_data)->>'confidence_score')::float) as avg_confidence,
            AVG((json_array_elements(tables_data)->>'data_quality_score')::float) as avg_quality
        FROM documents 
        WHERE tables_data IS NOT NULL
        GROUP BY json_array_elements(tables_data)->>'extraction_method'
    """)).fetchall()
    
    # Data quality distribution
    quality_stats = db.execute(text("""
        SELECT 
            CASE 
                WHEN (json_array_elements(tables_data)->>'data_quality_score')::float >= 0.8 THEN 'high'
                WHEN (json_array_elements(tables_data)->>'data_quality_score')::float >= 0.6 THEN 'medium'
                ELSE 'low'
            END as quality_level,
            COUNT(*) as count
        FROM documents 
        WHERE tables_data IS NOT NULL
        GROUP BY quality_level
    """)).fetchall()
    
    return {
        "document_statistics": {
            "total_documents": doc_stats.total_documents or 0,
            "documents_with_tables": doc_stats.docs_with_tables or 0,
            "total_tables_extracted": doc_stats.total_tables or 0,
            "average_tables_per_document": round(float(doc_stats.avg_tables_per_doc or 0), 2)
        },
        "table_type_distribution": {
            row.table_type or "unknown": row.count 
            for row in table_type_stats
        },
        "extraction_methods": {
            row.method or "unknown": {
                "count": row.count,
                "avg_confidence": round(float(row.avg_confidence or 0), 2),
                "avg_quality": round(float(row.avg_quality or 0), 2)
            }
            for row in extraction_method_stats
        },
        "data_quality_distribution": {
            row.quality_level: row.count 
            for row in quality_stats
        }
    }



# Tabular Data Processing Functions
def _apply_size_limits_to_task_result(result: dict) -> dict:
    """
    Apply size limits to async task results to prevent browser crashes.
    This ensures that async task status responses don't contain large datasets.
    """
    if not isinstance(result, dict):
        return result
    
    # Create a copy to avoid modifying the original
    limited_result = result.copy()
    
    # Limit table_preview data if present (for tabular files)
    if 'table_preview' in limited_result and isinstance(limited_result['table_preview'], list):
        preview_data = limited_result['table_preview']
        if len(preview_data) > config.large_file.max_response_rows:
            limited_result['table_preview'] = preview_data[:config.large_file.max_response_rows]
            limited_result['preview_truncated'] = True
            limited_result['preview_sample_size'] = config.large_file.max_response_rows
            limited_result['total_preview_rows'] = len(preview_data)
            logger.info(f"Task result preview truncated: showing {config.large_file.max_response_rows} of {len(preview_data)} rows")
        else:
            limited_result['preview_truncated'] = False
    
    # Limit any tables data if present (for regular documents)
    if 'tables' in limited_result and isinstance(limited_result['tables'], list):
        limited_tables = []
        for table in limited_result['tables']:
            if isinstance(table, dict):
                limited_table = table.copy()
                
                # Limit table rows/data
                if 'rows' in limited_table and isinstance(limited_table['rows'], list):
                    rows = limited_table['rows']
                    if len(rows) > config.large_file.max_response_rows:
                        limited_table['rows'] = rows[:config.large_file.max_response_rows]
                        limited_table['response_truncated'] = True
                        limited_table['response_sample_size'] = config.large_file.max_response_rows
                        limited_table['total_rows_available'] = len(rows)
                    else:
                        limited_table['response_truncated'] = False
                
                # Limit table data field (for CSV-style data)
                if 'data' in limited_table and isinstance(limited_table['data'], list):
                    data = limited_table['data']
                    if len(data) > config.large_file.max_response_rows:
                        limited_table['data'] = data[:config.large_file.max_response_rows]
                        limited_table['data_truncated'] = True
                        limited_table['data_sample_size'] = config.large_file.max_response_rows
                        limited_table['total_data_rows'] = len(data)
                    else:
                        limited_table['data_truncated'] = False
                
                limited_tables.append(limited_table)
            else:
                limited_tables.append(table)
        
        limited_result['tables'] = limited_tables
    
    # Add metadata about size limiting
    limited_result['size_limits_applied'] = True
    limited_result['max_response_rows'] = config.large_file.max_response_rows
    
    return limited_result

def _limit_table_data_for_response(table_data: dict, max_rows: int = None) -> dict:
    """Limit table data size for API responses to prevent browser crashes"""
    if max_rows is None:
        max_rows = config.large_file.max_response_rows
    
    # Create a copy to avoid modifying original data
    limited_data = table_data.copy()
    
    # If data exists and is too large, truncate it
    if 'data' in limited_data and isinstance(limited_data['data'], list):
        original_size = len(limited_data['data'])
        if original_size > max_rows:
            limited_data['data'] = limited_data['data'][:max_rows]
            limited_data['response_truncated'] = True
            limited_data['response_sample_size'] = max_rows
            limited_data['total_rows_available'] = original_size
            logger.warning(f"Response truncated: showing {max_rows} of {original_size} rows to prevent browser crash")
        else:
            limited_data['response_truncated'] = False
            limited_data['response_sample_size'] = original_size
    
    return limited_data

def _is_tabular_file(filename: str, content: bytes) -> bool:
    """Check if uploaded file is tabular format (CSV, Excel, TSV)"""
    from src.services.tabular_processor import TabularProcessor
    
    # First check: only consider files with explicit tabular extensions
    filename_lower = filename.lower()
    if filename_lower.endswith(('.csv', '.tsv', '.xlsx', '.xls')):
        return TabularProcessor.detect_file_type(filename, content) is not None
    
    # Don't treat code files or other structured text as tabular
    return False

async def _process_tabular_as_table(file: UploadFile, content: bytes, start_time: float, db: Session):
    """Process tabular file (CSV, Excel, TSV) as structured table data"""
    import hashlib
    from src.adapters.database.models import DocumentRecord
    from src.services.tabular_processor import TabularProcessor
    
    try:
        # Detect file type
        file_type = TabularProcessor.detect_file_type(file.filename, content)
        if not file_type:
            raise ValueError("Unable to detect tabular file type")
        
        # Load as DataFrame
        df = TabularProcessor.load_dataframe(content, file_type, file.filename)
        
        # Create table data structure
        table_data = TabularProcessor.create_table_data(df, file_type, file.filename)
        
        # Store in database
        file_hash = hashlib.sha256(content).hexdigest()
        existing = db.query(DocumentRecord).filter(DocumentRecord.file_hash == file_hash).first()
        
        if existing:
            # Update existing record
            existing.tables_data = [table_data]
            existing.table_count = 1
            action = "updated"
            document_id = existing.id
        else:
            # Create new record
            new_doc = DocumentRecord(
                filename=file.filename,
                file_extension=f".{file_type}",  # Set file extension
                file_size=len(content),  # Set file size in bytes
                file_hash=file_hash,
                full_text=f"{file_type.upper()} file with {len(df)} rows and {len(df.columns)} columns",
                page_count=1,
                has_ocr_content=0,  # Use integer 0 instead of boolean False
                processing_method=f"{file_type}_parser",
                table_count=1,
                tables_data=[table_data]
            )
            db.add(new_doc)
            db.commit()
            db.refresh(new_doc)
            action = "created"
            document_id = new_doc.id
        
        db.commit()
        
        # Calculate processing time
        processing_time = round((time.time() - start_time) * 1000)
        
        # Return table-formatted response
        return {
            "id": document_id,
            "filename": file.filename,
            "full_text": f"{file_type.upper()} table with {len(df)} rows and {len(df.columns)} columns",
            "page_count": 1,
            "has_ocr_content": 0,  # Use integer 0 instead of boolean False
            "processing_method": f"{file_type}_parser",
            "table_count": 1,
            "tables": [table_data],
            "action": action,
            "processing_time_ms": processing_time,
            "file_size_bytes": len(content),
            "data_format": "table",
            "table_preview": TabularProcessor.get_preview_data(df),
            "table_info": {
                "shape": f"{len(df)} rows × {len(df.columns)} columns",
                "columns": list(df.columns),
                "data_types": {col: str(df[col].dtype) for col in df.columns}
            },
            "data_quality": TabularProcessor.analyze_data_quality(df)
        }
        
    except Exception as e:
        logger.error(f"Tabular file processing failed: {e}")
        # Provide more helpful error messages for common issues
        error_msg = str(e)
        if "Expected" in error_msg and "fields" in error_msg:
            error_msg = f"CSV file has inconsistent number of columns. This often happens with malformed CSV files. Original error: {error_msg}"
        elif "tokenizing data" in error_msg.lower():
            error_msg = f"CSV file format is invalid or corrupted. Please check the file structure. Original error: {error_msg}"
        
        raise HTTPException(status_code=400, detail=f"Failed to process {file_type.upper() if 'file_type' in locals() else 'tabular'} file: {error_msg}")

@app.post("/extract/table/")
async def extract_tabular_data(file: UploadFile = File(...), db: Session = Depends(get_db)):
    """Dedicated endpoint for tabular data processing (CSV, Excel, TSV)"""
    import time
    start_time = time.time()
    
    content = await file.read()
    
    if not _is_tabular_file(file.filename, content):
        raise HTTPException(status_code=400, detail="File is not a valid tabular format (CSV, Excel, TSV)")
    
    return await _process_tabular_as_table(file, content, start_time, db)

@app.get("/health")
async def health_check():
    """Enhanced health check with table extraction status."""
    health = {
        "status": "healthy",
        "backend": "celery" if USE_CELERY else "background_tasks"
    }
    
    # Check table extraction capabilities
    try:
        import fitz
        health["pdf_table_extraction"] = "available"
    except ImportError:
        health["pdf_table_extraction"] = "unavailable"
        health["status"] = "degraded"
    
    try:
        import docx
        health["docx_table_extraction"] = "available"
    except ImportError:
        health["docx_table_extraction"] = "unavailable"
        health["status"] = "degraded"
    
    # Check tabular data processing
    try:
        import pandas
        health["tabular_data_processing"] = "available"
    except ImportError:
        health["tabular_data_processing"] = "unavailable"
        health["status"] = "degraded"
    
    try:
        import openpyxl
        health["excel_processing"] = "available"
    except ImportError:
        health["excel_processing"] = "unavailable"
    
    return health